# -*- coding: utf-8 -*-
"""Informe de investigación de la entidad en Word.

    python generar_informe.py --contenido informe.json \
        --evidencia-societario "....json" --evidencia-sector-publico "....json" \
        --evidencia-riesgos "....json" \
        --cliente "RAZON SOCIAL, S.A." --cif A00000000 --cierre 31/12/24 \
        --ventana "31/12/2019 - 26/08/2026" --generado 2026-08-26 --ref "" \
        --salida "<expediente>/InformesGesia/InvestigacionEntidad/....docx"

`--generado` es obligatorio: nada en este proyecto lee el reloj.

La cabecera lleva los mismos campos y en el mismo orden que los papeles de
continuidad de saldos e identificación de riesgos, para que los papeles del
expediente se lean igual.

La sección de fuentes NO la redacta el modelo: se construye aquí uniendo los
`fuentes_consultadas` de la evidencia JSON de los tres agentes, tal cual se
archivó. Así el informe no puede citar una fuente que la evidencia no respalde.
"""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

MARCAS = "✓⚠○"

# Limitaciones ESTRUCTURALES de las fuentes públicas gratuitas: son de las
# fuentes, no de la entidad, y se repiten en todos los informes. Texto fijo para
# que dos clientes distintos reciban la misma redacción; el modelo solo marca
# cuáles aplican por su id. Verificadas en la calibración con expediente real
# del 26/08/2026.
CATALOGO_LIMITACIONES = {
    "borme_sin_historico": (
        "No existe un buscador gratuito que indexe el histórico completo de actos "
        "del BORME por entidad; los actos citados proceden de búsquedas puntuales "
        "sobre boe.es y pueden no cubrir todos los publicados. Pueden existir "
        "actos intermedios no detectados, incluidos depósitos de cuentas."),
    "registro_mercantil_pago": (
        "El detalle del Registro Mercantil (situación vigente de cargos, depósito "
        "de cuentas, capital social) es de pago; la parte gratuita es escasa. Los "
        "huecos societarios de este informe son atribuibles a esa limitación, no "
        "a la entidad."),
    "aeat_censal": (
        "El estado censal en la AEAT exige certificado electrónico y no es "
        "consultable por terceros; no se ha verificado."),
    "place_sin_api": (
        "PLACE (contrataciondelestado.es) no ofrece un servicio público de "
        "consulta directa utilizable por herramientas automatizadas; la "
        "comprobación se hace por búsqueda web restringida al dominio, con un "
        "margen de incertidumbre residual mayor que el de BDNS."),
    "cendoj_no_fetchable": (
        "El buscador de CENDOJ no devuelve resultados legibles de forma fiable a "
        "herramientas automatizadas; la conclusión «sin datos» se apoya en "
        "búsqueda combinada, no en una exploración manual exhaustiva del buscador "
        "nativo."),
    "boletines_no_fetchable": (
        "Los buscadores de boletines autonómicos y provinciales presentan la "
        "misma limitación de lectura automatizada que el CENDOJ."),
    "prensa_no_indexada": (
        "La ausencia de piezas en cabeceras de prensa nacional se basa en los "
        "resultados del motor de búsqueda utilizado; no puede descartarse que "
        "existan piezas no indexadas."),
    "cnae_no_oficial": (
        "El CNAE no se publica en ninguna fuente registral gratuita; cuando "
        "consta, procede de agregadores (fiabilidad media)."),
}

DECLARACION = (
    "La información de este documento la ha recopilado un modelo de lenguaje en "
    "fuentes públicas de internet, con cada hecho citado a su fuente y fecha de "
    "consulta, y la evidencia archivada junto a este informe. No es una "
    "conclusión de auditoría: la evaluación del efecto de esta información en el "
    "encargo es del auditor, que es quien firma.")

LEYENDA = ("✓ hecho verificado en la fuente citada · ⚠ indicio, resultado del "
           "análisis cruzado · ○ sin información: consta que se buscó y no hay "
           "dato público")

# Orden de presentación de la tabla de fuentes: primero las que dieron datos
# (las «consultadas con éxito»), después las vacías y al final las inaccesibles.
ORDEN_RESULTADO = ["con datos", "sin datos", "inaccesible"]


def salida_utf8() -> None:
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    except AttributeError:
        pass


def parrafo(doc, texto: str, cursiva=False, tam=10, negrita=False):
    p = doc.add_paragraph()
    t = p.add_run(texto)
    t.italic, t.bold = cursiva, negrita
    t.font.size = Pt(tam)
    return p


def tabla(doc, cabeceras: list, filas: list, estilo="Light Grid Accent 1"):
    """Tabla con o sin fila de cabecera (sin cabecera si todas vienen vacías)."""
    con_cabecera = any(str(c).strip() for c in cabeceras)
    t = doc.add_table(rows=1 if con_cabecera else 0, cols=len(cabeceras))
    t.style = estilo
    if con_cabecera:
        for i, c in enumerate(cabeceras):
            celda = t.rows[0].cells[i]
            celda.text = str(c)
            for p in celda.paragraphs:
                for r in p.runs:
                    r.bold = True
    for fila in filas:
        celdas = t.add_row().cells
        for i, v in enumerate(fila):
            celdas[i].text = "" if v is None else str(v)
    return t


def validar_contenido(contenido: dict) -> list:
    """Devuelve la lista de defectos. Con defectos NO se genera informe."""
    defectos = []
    for nombre, items in (contenido.get("secciones") or {}).items():
        for item in items:
            if item.get("marca") not in tuple(MARCAS):
                defectos.append("la sección '%s' lleva un item con marca '%s'; "
                                "las válidas son ✓ ⚠ ○"
                                % (nombre, item.get("marca")))
            if item.get("marca") == "✓" and not item.get("fuente"):
                defectos.append("hecho ✓ sin fuente en la sección '%s': «%s…»"
                                % (nombre, str(item.get("texto", ""))[:60]))
    for clave in contenido.get("limitaciones_catalogo", []):
        if clave not in CATALOGO_LIMITACIONES:
            defectos.append("id de limitación desconocido: '%s'. No se inventan "
                            "limitaciones de catálogo" % clave)
    if not contenido.get("limitaciones_catalogo") and \
            not contenido.get("limitaciones_especificas"):
        defectos.append("el apartado de limitaciones no puede quedar vacío")
    return defectos


def items_seccion(doc, items: list):
    """Escribe los items de una sección con su marca y su fuente."""
    if not items:
        parrafo(doc, "○ Sin contenido en esta sección.", tam=10)
        return
    for item in items:
        texto = item["marca"] + " " + item["texto"]
        if item.get("fuente"):
            texto += " (" + item["fuente"] + ")"
        if item.get("incertidumbre"):
            texto += " — incertidumbre " + item["incertidumbre"]
        parrafo(doc, texto, tam=10)


def fuentes_de_evidencia(rutas: dict) -> list:
    """Une los fuentes_consultadas de la evidencia. Fila: bloque, nombre, url,
    fecha, resultado. Una evidencia ausente (agente degradado) se anota."""
    filas = []
    for bloque, ruta in rutas.items():
        if not ruta:
            filas.append((bloque, "— agente degradado: sin evidencia —",
                          "", "", "inaccesible"))
            continue
        datos = json.loads(Path(ruta).read_text(encoding="utf-8"))
        for f in datos.get("fuentes_consultadas", []):
            filas.append((bloque, f.get("nombre", "—"), f.get("url", ""),
                          f.get("fecha_consulta", ""), f.get("resultado", "—")))

    def clave_orden(fila):
        resultado = fila[4]
        for i, prefijo in enumerate(ORDEN_RESULTADO):
            if str(resultado).startswith(prefijo):
                return i
        return len(ORDEN_RESULTADO)

    return sorted(filas, key=clave_orden)


def main() -> int:
    salida_utf8()
    p = argparse.ArgumentParser()
    p.add_argument("--contenido", required=True,
                   help="JSON con las secciones escritas por el modelo")
    p.add_argument("--evidencia-societario")
    p.add_argument("--evidencia-sector-publico")
    p.add_argument("--evidencia-riesgos")
    p.add_argument("--cliente", required=True)
    p.add_argument("--cif", required=True)
    p.add_argument("--cierre", required=True)
    p.add_argument("--ventana", required=True,
                   help="ventana temporal investigada, p. ej. '31/12/2019 - 26/08/2026'")
    p.add_argument("--generado", required=True)
    # La referencia se PREGUNTA al usuario. Si no tiene una, se pasa vacía y la
    # celda queda en blanco: nunca se inventa un código que no existe en Gesia.
    p.add_argument("--ref", required=True,
                   help="referencia del papel; cadena vacía si el usuario no la tiene")
    p.add_argument("--salida", required=True)
    args = p.parse_args()

    contenido = json.loads(Path(args.contenido).read_text(encoding="utf-8"))
    defectos = validar_contenido(contenido)
    if defectos:
        for d in defectos:
            print("DEFECTO: " + d)
        return 2

    secciones = contenido.get("secciones") or {}
    evidencias = {
        "Societario": args.evidencia_societario,
        "Sector público": args.evidencia_sector_publico,
        "Riesgos legales": args.evidencia_riesgos,
    }

    doc = Document()
    doc.core_properties.title = "Investigación de la entidad en fuentes públicas"

    # ── Cabecera común de los papeles del expediente ──────────────────────────
    parrafo(doc, args.cliente, negrita=True, tam=12)
    parrafo(doc, "AUDITORIA A " + args.cierre)
    h = doc.add_heading("INVESTIGACIÓN DE LA ENTIDAD EN FUENTES PÚBLICAS", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    parrafo(doc, "Fuente: fuentes públicas de internet (BORME/BOE, PLACE, BDNS, "
                 "CENDOJ, boletines y prensa); evidencia JSON archivada en la "
                 "subcarpeta «evidencia» junto a este informe")
    tabla(doc, ["", ""], [["REF. P.T.", args.ref],
                          ["Realizado", ""],
                          ["Verificado", ""]])
    parrafo(doc, "Conocimiento de la entidad y su entorno por fuentes externas · "
                 "NIA-ES 315", cursiva=True, tam=9)
    parrafo(doc, "Generado: " + args.generado + " · Ventana investigada: "
                 + args.ventana, tam=9)
    parrafo(doc, LEYENDA, cursiva=True, tam=9)
    parrafo(doc, DECLARACION, cursiva=True, tam=9)

    # ── 1. Identificación ─────────────────────────────────────────────────────
    doc.add_heading("1. Identificación de la entidad", level=2)
    filas = [["Denominación", args.cliente], ["CIF", args.cif],
             ["Cierre del ejercicio auditado", args.cierre]]
    filas += [[i.get("campo", "—"), i.get("valor", "—")]
              for i in contenido.get("identificacion", [])]
    tabla(doc, ["", ""], filas)

    # ── 2 a 5: los bloques de hechos ──────────────────────────────────────────
    doc.add_heading("2. Estructura societaria", level=2)
    items_seccion(doc, secciones.get("societaria", []))
    doc.add_heading("3. Actividad y contexto sectorial", level=2)
    items_seccion(doc, secciones.get("actividad", []))
    doc.add_heading("4. Relación con el sector público", level=2)
    items_seccion(doc, secciones.get("sector_publico", []))
    doc.add_heading("5. Riesgos legales y reputacionales", level=2)
    items_seccion(doc, secciones.get("riesgos_legales", []))

    # ── 6. Análisis ───────────────────────────────────────────────────────────
    doc.add_heading("6. Análisis — elementos a contrastar", level=2)
    analisis = secciones.get("analisis", [])
    items_seccion(doc, analisis)
    if analisis:
        parrafo(doc, "Los elementos anteriores cruzan hechos ya citados; no son "
                     "conclusiones. Incertidumbre baja = respaldado por dos o más "
                     "fuentes; alta = hipótesis a partir de una coincidencia.",
                cursiva=True, tam=9)

    # ── 7. Limitaciones (obligatorio, y no puede quedar vacío) ────────────────
    doc.add_heading("7. Limitaciones del análisis", level=2)
    for clave in contenido.get("limitaciones_catalogo", []):
        parrafo(doc, "· " + CATALOGO_LIMITACIONES[clave], tam=10)
    for texto in contenido.get("limitaciones_especificas", []):
        parrafo(doc, "· " + texto, tam=10)

    # ── 8. Fuentes consultadas, desde la evidencia y no desde el modelo ──────
    doc.add_heading("8. Fuentes consultadas", level=2)
    filas_fuentes = fuentes_de_evidencia(evidencias)
    con_exito = sum(1 for f in filas_fuentes if str(f[4]).startswith("con datos"))
    parrafo(doc, str(len(filas_fuentes)) + " fuentes consultadas, "
                 + str(con_exito) + " con datos. Primero las que aportaron "
                 "información, después las consultadas sin resultado y las "
                 "inaccesibles: todas acreditan dónde se buscó.", tam=10)
    tabla(doc, ["Bloque", "Fuente", "URL", "Fecha consulta", "Resultado"],
          [list(f) for f in filas_fuentes])

    destino = Path(args.salida)
    destino.parent.mkdir(parents=True, exist_ok=True)
    doc.save(destino)
    print("Informe escrito en " + str(destino))
    print("  secciones: " + ", ".join(sorted(secciones)))
    print("  fuentes: " + str(len(filas_fuentes)) + " (" + str(con_exito)
          + " con datos)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
