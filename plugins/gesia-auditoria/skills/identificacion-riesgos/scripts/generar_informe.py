"""Informe de identificacion de riesgos en Word.

    python generar_informe.py --analisis analisis.json --seleccion seleccion.json \
        --cliente "RAZON SOCIAL, S.A." --cierre 31/12/25 --generado 2026-08-24 \
        --salida "<expediente>/InformesGesia/IdentificacionRiesgos/....docx"

`--generado` es obligatorio: nada en este proyecto lee el reloj.

Los graficos se dibujan con matplotlib si esta disponible y, si no, el informe
sale con las mismas cifras en tabla. Un informe sin grafico se lee; un informe que
no se genera, no.
"""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from lib_riesgos import (
    CONCENTRACION,
    EJERCICIOS_PERDIDAS,
    OBLIGATORIOS,
    VARIACION,
    cargar_catalogo,
    riesgo_legible,
    salida_utf8,
)

MAX_HECHOS = 12

# Cuantos ratios de la revision analitica se listan en el informe.
MAX_RA = 12

DECLARACION = (
    "Los riesgos de este documento los ha propuesto un modelo de lenguaje a partir "
    "del catálogo de riesgos del máster de Gesia y de las cifras del expediente. "
    "No son una conclusión de auditoría: la valoración del riesgo, su relevancia "
    "para el encargo y los procedimientos que finalmente se apliquen son del "
    "auditor, que es quien firma."
)


def euro(v) -> str:
    if v is None:
        return "—"
    return format(float(v), ",.0f").replace(",", " ")


def porcentaje(v) -> str:
    if v is None:
        return "s/base"
    return format(abs(float(v)) * 100, ".1f") + " %"


def parrafo(doc, texto: str, cursiva=False, tam=10, negrita=False):
    p = doc.add_paragraph()
    t = p.add_run(texto)
    t.italic, t.bold = cursiva, negrita
    t.font.size = Pt(tam)
    return p


def texto_señales(señales: list) -> str:
    """Las señales medidas, en palabras, para la tabla del informe."""
    partes = []
    for s in señales:
        if s["tipo"] == "tendencia":
            partes.append(s["sentido"] + " " + str(s["años"]) + " años")
        elif s["tipo"] == "salto":
            veces = s.get("veces_lo_normal")
            partes.append("salto brusco"
                          + (" (x" + str(veces) + " lo habitual)" if veces else ""))
        elif s["tipo"] == "negativo_persistente":
            partes.append("negativo " + str(s["años"]) + " años")
        elif s["tipo"] == "cambio_de_signo":
            partes.append("cambia de signo")
    return " · ".join(partes)


def tabla(doc, cabeceras: list, filas: list, estilo="Light Grid Accent 1"):
    """Tabla con o sin fila de cabecera.

    Si todas las cabeceras vienen vacias es una tabla de pares clave-valor y NO
    se crea fila de cabecera: dejaba una fila en blanco arriba.
    """
    con_cabecera = any(str(c).strip() for c in cabeceras)
    t = doc.add_table(rows=1 if con_cabecera else 0, cols=len(cabeceras))
    t.style = estilo
    if con_cabecera:
        for i, c in enumerate(cabeceras):
            celda = t.rows[0].cells[i]
            celda.text = str(c)
            for p in celda.paragraphs:
                for r in p.runs:
                    r.bold = True
    for fila in filas:
        celdas = t.add_row().cells
        for i, v in enumerate(fila):
            celdas[i].text = "" if v is None else str(v)
    return t


def grafico_resultado(res: dict, destino: Path):
    """Evolucion del resultado en los cinco ejercicios. Devuelve la ruta o None."""
    if not res.get("encontrado"):
        return None
    try:
        import matplotlib
        matplotlib.use("Agg")               # sin pantalla: se escribe a fichero
        import matplotlib.pyplot as plt
    except Exception:                        # noqa: BLE001
        return None

    serie = list(reversed(res["serie"]))
    etiquetas = list(reversed(res.get("ejercicios") or
                              [str(i) for i in range(len(serie))]))
    fig, ax = plt.subplots(figsize=(6.2, 2.6))
    colores = ["#c0392b" if v < 0 else "#2c6e91" for v in serie]
    ax.bar(etiquetas, serie, color=colores)
    ax.axhline(0, color="#444", linewidth=0.8)
    ax.set_title("Resultado del ejercicio", fontsize=9)
    ax.tick_params(labelsize=8)
    ax.spines[["top", "right"]].set_visible(False)
    ax.yaxis.set_major_formatter(
        lambda v, _p: format(v / 1000, ",.0f").replace(",", " ") + " k")
    fig.tight_layout()
    ruta = destino / "_grafico_resultado.png"
    fig.savefig(ruta, dpi=160)
    plt.close(fig)
    return ruta


def main() -> int:
    salida_utf8()
    p = argparse.ArgumentParser()
    p.add_argument("--analisis", required=True)
    p.add_argument("--seleccion", required=True)
    p.add_argument("--analisis-ra", help="salida de analizar_ra.py, si la hay")
    # El analisis escrito lo aporta el modelo: las tablas dicen QUE pasa y este
    # texto dice que significa. Sin el, el informe salta de las cifras a los
    # riesgos y el revisor no ve el razonamiento intermedio.
    p.add_argument("--narrativa", help="JSON con los análisis escritos por bloque")
    p.add_argument("--cliente", default="")
    p.add_argument("--cierre", default="")
    p.add_argument("--sector", default="")
    p.add_argument("--generado", required=True)
    # La referencia se PREGUNTA al usuario. No hay valor por defecto: inventar
    # una referencia de papel de trabajo es meter en el expediente un codigo
    # que no existe en el indice de Gesia.
    p.add_argument("--ref", required=True,
                   help="referencia del papel, la que diga el usuario")
    # Sin --fuente se compone con lo que REALMENTE se ha usado, y con la
    # denominacion que corresponda a la entidad: en una sin animo de lucro no se
    # puede escribir "PyG".
    p.add_argument("--fuente", default=None,
                   help="texto de la línea Fuente; si se omite, se compone solo")
    p.add_argument("--catalogo")
    p.add_argument("--salida", required=True)
    args = p.parse_args()

    analisis = json.loads(Path(args.analisis).read_text(encoding="utf-8"))
    seleccion = json.loads(Path(args.seleccion).read_text(encoding="utf-8"))
    ra = (json.loads(Path(args.analisis_ra).read_text(encoding="utf-8"))
          if args.analisis_ra else None)
    narrativa = (json.loads(Path(args.narrativa).read_text(encoding="utf-8"))
                 if args.narrativa else {})
    elegidos = seleccion.get("riesgos") if isinstance(seleccion, dict) else seleccion
    catalogo = cargar_catalogo(args.catalogo)
    por_id = catalogo["por_id"]

    entidad = analisis.get("entidad", {})
    como_se_llama = entidad.get("denominacion_resultados", "cuenta de pérdidas y ganancias")

    doc = Document()
    doc.core_properties.title = "Identificación de riesgos de auditoría"

    # Los PNG de los graficos se incrustan en el .docx, asi que el fichero
    # suelto sobra en cuanto se guarda. Antes se avisaba al usuario de que lo
    # borrara, y es justo el paso que se olvida: queda basura en la carpeta del
    # expediente.
    temporales: list = []

    # ── Cabecera ──────────────────────────────────────────────────────────────
    # Los mismos campos y en el mismo orden que el papel de continuidad de saldos,
    # para que los papeles del expediente se lean igual: cliente, «AUDITORIA A
    # <fecha>», titulo, fuente, y el bloque REF. P.T. / Realizado / Verificado.
    parrafo(doc, args.cliente or "(cliente no indicado)", negrita=True, tam=12)
    parrafo(doc, "AUDITORIA A " + (args.cierre or "—"))
    h = doc.add_heading("IDENTIFICACIÓN DE RIESGOS DE AUDITORÍA", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if args.fuente:
        fuente = args.fuente
    else:
        partes = ["Balance de situación", como_se_llama.capitalize()]
        if ra and ra.get("con_señal"):
            partes.append("módulo de Revisión analítica")
        fuente = (", ".join(partes[:-1]) + " y " + partes[-1]
                  + " plurianuales disponibles en Gesia")
    parrafo(doc, "Fuente: " + fuente)

    # Realizado y Verificado van en blanco: los rellena el auditor.
    tabla(doc, ["", ""], [["REF. P.T.", args.ref],
                          ["Realizado", ""],
                          ["Verificado", ""]])

    parrafo(doc, "NIA-ES 315 · Identificación y valoración de los riesgos de "
                 "incorrección material", cursiva=True, tam=9)

    # La materialidad y de donde sale. Una IR_T calculada en Gesia y una dada a
    # mano en el momento no acreditan lo mismo, y el papel tiene que distinguirlo.
    mat = analisis.get("materialidad") or {}
    if mat.get("aplicada"):
        origen = {"gesia": "Gesia", "manual": "Manual"}.get(mat.get("origen"),
                                                            "no indicado")
        texto_ir = (format(float(mat["ir_t"]), ",.2f").replace(",", " ")
                    + " EUR · Origen: " + origen)
    else:
        texto_ir = "SIN APLICAR — se han analizado todos los epígrafes"

    tabla(doc, ["", ""], [
        ["Importancia relativa (IR_T)", texto_ir],
        ["Sector", args.sector or "no indicado"],
        ["Tipo de entidad", "Sin ánimo de lucro"
         if entidad.get("sin_animo_de_lucro") else "Sociedad mercantil"],
        ["Riesgos identificados", str(len(elegidos))],
        ["Catálogo de riesgos", catalogo.get("version", "")],
        ["Generado el", args.generado],
    ])

    if entidad.get("sin_animo_de_lucro"):
        parrafo(doc, "La entidad se ha tratado como entidad sin ánimo de lucro por "
                     "el vocabulario de sus cuentas anuales (" +
                     ", ".join(entidad.get("terminos_detectados", [])[:4]) +
                     "). En este informe se habla de «" + como_se_llama + "».",
                cursiva=True, tam=9)

    # ── Cómo se ha hecho ──────────────────────────────────────────────────────
    doc.add_heading("Alcance y método", level=1)
    parrafo(doc, "Se han analizado el balance y la " + como_se_llama + " de los "
                 "cinco últimos ejercicios registrados en el expediente. Sobre "
                 "esas cifras se han aplicado tres umbrales: variaciones "
                 "interanuales iguales o superiores al " + str(int(VARIACION * 100))
                 + " %, epígrafes que concentran el " + str(int(CONCENTRACION * 100))
                 + " % o más de su masa patrimonial, y pérdidas en "
                 + str(EJERCICIOS_PERDIDAS) + " o más de los cinco ejercicios.")
    parrafo(doc, "Los riesgos se han seleccionado del catálogo de riesgos del máster "
                 "de Gesia; no se ha redactado ninguno nuevo. Los riesgos "
                 + " y ".join(str(i) for i in OBLIGATORIOS) + " se incluyen siempre, "
                 "con independencia de las cifras.")

    # ── Lo que dicen las cifras ───────────────────────────────────────────────
    doc.add_heading("Hechos observados", level=1)
    res = analisis.get("resultado_ejercicio", {})
    if res.get("encontrado"):
        doc.add_heading("Resultado de los cinco ejercicios", level=2)
        ejer = res.get("ejercicios") or []
        tabla(doc, ejer or [""] * len(res["serie"]), [[euro(v) for v in res["serie"]]])
        # El auditor lee esta tabla sin saber de donde sale, asi que hay que
        # decirle en que signo esta: en Gesia el beneficio es acreedor y sale
        # negativo, y aqui va al contrario.
        if res.get("signo_invertido"):
            parrafo(doc, "Beneficio en positivo y pérdida en negativo. En la "
                         "cuenta de pérdidas y ganancias de Gesia el resultado "
                         "figura con signo contable —el beneficio, acreedor— y "
                         "aquí se presenta invertido.", cursiva=True)
        if res.get("supera_umbral"):
            parrafo(doc, "Pérdidas en " + str(res["ejercicios_con_perdidas"])
                    + " de los cinco ejercicios.", negrita=True)
        grafico = grafico_resultado(res, Path(args.salida).parent)
        if grafico:
            doc.add_picture(str(grafico), width=Cm(15.5))
            temporales.append(grafico)

    for clave, titulo in (("balance", "Balance"),
                          ("resultados", como_se_llama.capitalize())):
        bloque = analisis.get(clave, {})
        var = bloque.get("variaciones", [])[:MAX_HECHOS]
        if var:
            doc.add_heading("Variaciones relevantes · " + titulo, level=2)
            tabla(doc,
                  ["Epígrafe", "Concepto", "Ejercicio", "Anterior", "Cambio"],
                  # Se pinta el importe de presentacion, no el saldo crudo: la
                  # variacion se mide en magnitud y una tabla que dijera
                  # «-492.657 ... aumenta un 430,7 %» se contradice a si misma.
                  # El dato de Gesia queda en el JSON, para poder contrastarlo.
                  [[v["epigrafe"],
                    v["concepto"] + (" (subtotal)" if v.get("es_subtotal") else ""),
                    euro(v.get("actual_presentado", v["actual"])),
                    euro(v.get("anterior_presentado", v["anterior"])),
                    v["sentido"] + " " + porcentaje(v.get("variacion_absoluta"))]
                   for v in var])
            if any(v.get("presentacion") == "magnitud" for v in var):
                parrafo(doc, "Importes en magnitud, sin el signo contable de "
                             "Gesia: es lo que concuerda con la columna de "
                             "cambio, que se mide en valor absoluto. Cuando un "
                             "epígrafe cambia de lado se dice «cambia de signo» "
                             "y ahí sí se presenta con su signo —beneficio en "
                             "positivo—.", cursiva=True)
        con = bloque.get("concentraciones", [])[:MAX_HECHOS]
        if con:
            doc.add_heading("Concentración · " + titulo, level=2)
            tabla(doc, ["Epígrafe", "Concepto", "Importe", "Peso"],
                  # El peso ya se calcula en valor absoluto, asi que el importe
                  # se pinta igual: en magnitud.
                  [[c["epigrafe"], c["concepto"], euro(abs(c["importe"])),
                    porcentaje(c["peso"])] for c in con])
        if narrativa.get(clave):
            doc.add_heading("Análisis · " + titulo, level=2)
            for p_ in str(narrativa[clave]).split("\n\n"):
                if p_.strip():
                    parrafo(doc, p_.strip())

    # ── Los riesgos ───────────────────────────────────────────────────────────
    if ra and ra.get("con_señal"):
        doc.add_heading("Revisión analítica", level=2)
        parrafo(doc, "De los " + str(ra["n_elementos"]) + " ratios activos en el "
                     "expediente, " + str(ra["n_con_señal"]) + " presentan algún "
                     "comportamiento reseñable. Se recogen los más marcados.")
        filas_ra = []
        for e in ra["con_señal"][:MAX_RA]:
            valor = None
            if ra.get("ejercicios"):
                valor = e["valores"].get(ra["ejercicios"][0])
            filas_ra.append([
                e["codigo"] + " · " + e["elemento"],
                e["familia"],
                (format(valor, ",.4f").replace(",", " ") if valor is not None else "—"),
                texto_señales(e["señales"]),
            ])
        tabla(doc, ["Ratio", "Familia", "Ejercicio", "Comportamiento"], filas_ra)
        if narrativa.get("revision_analitica"):
            for p_ in str(narrativa["revision_analitica"]).split("\n\n"):
                if p_.strip():
                    parrafo(doc, p_.strip())

    doc.add_heading("Riesgos identificados", level=1)
    # El area, el tipo y la valoracion salen del CATALOGO, no de la seleccion: la
    # seleccion solo aporta el id, la justificacion y la evidencia.
    resumen = []
    for r in elegidos:
        i = int(r["id"])
        if i not in por_id:
            continue
        leg = riesgo_legible(por_id[i])
        # En el resumen va SOLO el riesgo de incorreccion material: es el que
        # resume a los otros dos. Los tres niveles se ven en el detalle.
        resumen.append([i, leg["area"], leg["nombre"], leg["tipo"],
                        leg["calificacion"] or "—",
                        leg["riesgo_incorreccion"] or "—"])
    tabla(doc, ["Id", "Área", "Riesgo", "Tipo", "Calificación",
                "Incorrección material"], resumen)

    for r in elegidos:
        i = int(r["id"])
        if i not in por_id:
            continue
        leg = riesgo_legible(por_id[i], con_procedimientos=True)
        doc.add_heading(str(i) + " · " + leg["nombre"], level=2)

        detalle = [
            ("Área", leg["area"] + (" — " + leg["area_nombre"] if leg["area_nombre"] else "")),
            ("Tipo de riesgo", leg["tipo"]),
            ("Significativo", "Sí" if leg["significativo"] else "No"),
            ("Afirmaciones", ", ".join(leg["aserciones"]) or "—"),
            ("Calificación del riesgo", leg["calificacion"] or "—"),
            ("Riesgo inherente", leg["riesgo_inherente"] or "—"),
            ("Riesgo de control interno", leg["riesgo_control"] or "—"),
            ("Riesgo de incorrección material", leg["riesgo_incorreccion"] or "—"),
            ("Referencia", leg["referencia"] or "—"),
        ]
        ev = r.get("evidencia") or {}
        if i in OBLIGATORIOS and not ev:
            detalle.append(("Motivo de inclusión",
                            "Riesgo de inclusión obligatoria en todo encargo"))
        elif ev.get("ratio") or ev.get("elemento"):
            # Evidencia de la revision analitica: un ratio que se comporta mal
            # justifica un riesgo igual que una variacion en un epigrafe.
            detalle.append(("Ratio", (str(ev.get("ratio", "")) + " "
                                      + str(ev.get("elemento", ""))).strip()))
            v = ev.get("valor", ev.get("importe"))
            detalle.append(("Valor",
                            format(float(v), ",.4f").replace(",", " ")
                            if v is not None else "—"))
            if ev.get("señal"):
                detalle.append(("Comportamiento", str(ev["señal"])))
        else:
            detalle.append(("Epígrafe", (ev.get("epigrafe", "") + " "
                                         + ev.get("concepto", "")).strip() or "—"))
            detalle.append(("Importe", euro(ev.get("importe"))))
            detalle.append(("Variación", porcentaje(ev.get("variacion"))
                            if ev.get("variacion") is not None else "—"))
        tabla(doc, ["", ""], [[k, v] for k, v in detalle])

        if r.get("justificacion"):
            parrafo(doc, str(r["justificacion"]))
        if leg["descripcion"]:
            doc.add_heading("Descripción del riesgo", level=3)
            for bloque in leg["descripcion"].split("\n\n"):
                if bloque.strip():
                    parrafo(doc, bloque.strip(), tam=9)
        if leg.get("procedimientos"):
            doc.add_heading("Procedimientos previstos", level=3)
            for linea in leg["procedimientos"].split("\n"):
                if linea.strip():
                    parrafo(doc, linea.strip(), tam=9)

    # ── Cierre ────────────────────────────────────────────────────────────────
    if narrativa.get("resumen"):
        doc.add_heading("Resumen", level=1)
        for p_ in str(narrativa["resumen"]).split("\n\n"):
            if p_.strip():
                parrafo(doc, p_.strip())

    doc.add_heading("Nota", level=1)
    parrafo(doc, DECLARACION, cursiva=True, tam=9)

    salida = Path(args.salida)
    salida.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(salida)
    except PermissionError:
        print("ERROR: no se puede escribir " + str(salida))
        print("       El fichero está abierto en Word. Ciérralo y repite.")
        return 2

    print("Escrito: " + str(salida))
    print("  " + str(len(elegidos)) + " riesgos")
    hubo_grafico = bool(temporales)
    for f in temporales:
        try:
            f.unlink()
        except OSError:
            pass                    # si no se deja borrar, el informe ya está bien
    print("  gráficos: " + ("sí" if hubo_grafico
                            else "no (matplotlib no disponible; las cifras van en tabla)"))
    return 0


if __name__ == "__main__":
    sys.exit(main())
