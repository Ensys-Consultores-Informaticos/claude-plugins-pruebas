"""Papel de trabajo en Word: que se respondio, por que y con que evidencia.

El .xlsx no tiene sitio para justificar nada — la columna `Comentario` trae la
referencia normativa del master y sobrescribirla destruye dato. Sin este
documento el criterio no consta en ninguna parte, y un papel de trabajo cuyo
criterio no consta es indefendible en una inspeccion.

    python generar_papel.py --cuestionario trabajo/cuestionario.json \
        --respuestas trabajo/respuestas.json --generado 2026-08-23 \
        --cliente "RAZON SOCIAL, S.A." --cierre "31/12/25" --salida papel.docx

`--generado` es obligatorio a proposito: nada en este proyecto lee el reloj. Un
papel de trabajo tiene que poder regenerarse identico dentro de dos años.
"""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from lib_cuestionario import (
    RESPUESTAS,
    es_cabecera,
    leer_cuestionario,
    partir_en_secciones,
    salida_utf8,
)

# Cuantas preguntas se detallan por seccion en la tabla larga. Con 137 preguntas
# el documento ya son ~20 paginas; el tope evita que un cuestionario de 481
# (memoria consolidada) lo convierta en algo que nadie abre.
MAX_POR_SECCION = 40

DECLARACION = (
    "Las respuestas de este cuestionario las ha propuesto un modelo de lenguaje a partir "
    "de las cuentas anuales aportadas. NO son una conclusion de auditoria: son una "
    "propuesta que el auditor revisa, corrige y asume antes de importarla al expediente. "
    "Ante cualquier duda sobre si un desglose figura o no en la memoria, la respuesta "
    "emitida es 4 (Pendiente), nunca 1 (Si)."
)

LIMITACIONES = [
    "Solo se ha revisado si el desglose figura en la memoria. No se juzga si su contenido "
    "es correcto, ni si las cifras que contiene cuadran con la contabilidad.",
    "No se han contrastado las respuestas contra los saldos del expediente. Una pregunta "
    "respondida «no aplica» sobre un epigrafe en el que el cliente tiene saldo no se "
    "detecta en esta version.",
    "El balance y la cuenta de perdidas y ganancias no se han cuadrado contra el PDF.",
    "Las respuestas dependen de lo que el documento aportado permita leer. Un desglose "
    "presente pero redactado de forma que no se localice se responde 4 (Pendiente).",
]


def parrafo(doc, texto: str, cursiva: bool = False, tam: int = 10):
    p = doc.add_paragraph()
    t = p.add_run(texto)
    t.italic = cursiva
    t.font.size = Pt(tam)
    return p


def main() -> int:
    salida_utf8()
    p = argparse.ArgumentParser(description=__doc__)
    p.add_argument("--cuestionario", required=True)
    p.add_argument("--respuestas", required=True)
    p.add_argument("--generado", required=True, help="Fecha de generacion, AAAA-MM-DD")
    p.add_argument("--salida", required=True)
    # El cliente y el cierre los pasa quien llama, con lo que le dio
    # contexto_expediente. Antes se leian del API de Gesia, y desde Cowork eso
    # no responde: el papel salia con la cabecera en blanco y sin avisar de
    # nada, porque datos_encargo se comia la excepcion. Visto en cliente el
    # 23/08/2026.
    p.add_argument("--cliente", default="", help="razon social, de contexto_expediente")
    p.add_argument("--cierre", default="", help="fecha de cierre, de contexto_expediente")
    args = p.parse_args()

    cuestionario = leer_cuestionario(args.cuestionario)
    crudo = json.loads(Path(args.respuestas).read_text(encoding="utf-8"))
    lista = crudo["respuestas"] if isinstance(crudo, dict) else crudo
    por_orden = {str(r["orden"]): r for r in lista}

    cliente, cierre = args.cliente.strip(), args.cierre.strip()
    filas = cuestionario["filas"]
    secciones = partir_en_secciones(filas)
    preguntas = [f for f in filas if not es_cabecera(f)]

    doc = Document()
    doc.add_heading("Revisión del contenido de la memoria", level=0)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sub.add_run(cliente or "(cliente no disponible)").bold = True

    ficha = doc.add_table(rows=0, cols=2)
    ficha.style = "Light Grid Accent 1"
    for etiqueta, valor in [
        ("Cierre del ejercicio", cierre or "—"),
        ("Cuestionario", cuestionario["guia"]),
        ("Preguntas", str(len(preguntas))),
        ("Generado el", args.generado),
    ]:
        celda = ficha.add_row().cells
        celda[0].text = etiqueta
        celda[1].text = valor

    doc.add_heading("Cómo leer este documento", level=1)
    parrafo(doc, DECLARACION, cursiva=True)

    # Resumen.
    doc.add_heading("Resumen", level=1)
    reparto = {c: 0 for c in RESPUESTAS}
    for r in lista:
        reparto[int(r["respuesta"])] = reparto.get(int(r["respuesta"]), 0) + 1
    tabla = doc.add_table(rows=1, cols=3)
    tabla.style = "Light Grid Accent 1"
    for i, texto in enumerate(("Código", "Significado", "Preguntas")):
        tabla.rows[0].cells[i].text = texto
    for codigo in sorted(RESPUESTAS):
        fila = tabla.add_row().cells
        fila[0].text = str(codigo)
        fila[1].text = RESPUESTAS[codigo]
        fila[2].text = str(reparto.get(codigo, 0))

    # Lo que el auditor tiene que mirar primero.
    for codigo, titulo, entradilla in [
        (2, "Desgloses que no figuran en la memoria",
         "Preguntas respondidas «No». Son las que exigen decidir si falta información "
         "en la memoria o si la pregunta no era aplicable."),
        (4, "Pendientes de resolver",
         "Preguntas que el documento aportado no permite responder. Requieren "
         "intervención del auditor antes de dar el cuestionario por cerrado."),
    ]:
        seleccion = [r for r in lista if int(r["respuesta"]) == codigo]
        doc.add_heading(f"{titulo} ({len(seleccion)})", level=1)
        parrafo(doc, entradilla, cursiva=True, tam=9)
        if not seleccion:
            parrafo(doc, "Ninguna.")
            continue
        for r in seleccion:
            fila = next((f for f in preguntas if str(f["CodigoOrden"]) == str(r["orden"])), None)
            if fila is None:
                continue
            enc = doc.add_paragraph()
            enc.add_run(f"{fila['Punto']}.{fila['DesglosePunto']}  ").bold = True
            enc.add_run(" ".join(str(fila["Descripcion"]).split()))
            if r.get("motivo"):
                parrafo(doc, f"Motivo: {r['motivo']}", tam=9)
            if r.get("evidencia"):
                parrafo(doc, f"Evidencia: {r['evidencia']}", tam=9)

    # Detalle completo.
    doc.add_heading("Detalle por secciones", level=1)
    for seccion in secciones:
        if not seccion["preguntas"]:
            continue
        doc.add_heading(
            f"Punto {seccion['punto']} — {' '.join(str(seccion['titulo']).split())}", level=2
        )
        tabla = doc.add_table(rows=1, cols=4)
        tabla.style = "Light Grid Accent 1"
        for i, texto in enumerate(("#", "Pregunta", "Resp.", "Motivo y evidencia")):
            tabla.rows[0].cells[i].text = texto
        for fila in seccion["preguntas"][:MAX_POR_SECCION]:
            r = por_orden.get(str(fila["CodigoOrden"]), {})
            codigo = int(r["respuesta"]) if r.get("respuesta") is not None else None
            celdas = tabla.add_row().cells
            celdas[0].text = f"{fila['Punto']}.{fila['DesglosePunto']}"
            celdas[1].text = " ".join(str(fila["Descripcion"]).split())
            celdas[2].text = f"{codigo} {RESPUESTAS[codigo]}" if codigo in RESPUESTAS else "—"
            justificacion = " ".join(x for x in (r.get("motivo"), r.get("evidencia")) if x)
            celdas[3].text = justificacion
        sobran = len(seccion["preguntas"]) - MAX_POR_SECCION
        if sobran > 0:
            parrafo(doc, f"({sobran} preguntas más de esta sección, en el .xlsx)", cursiva=True, tam=9)

    doc.add_heading("Alcance y limitaciones", level=1)
    for texto in LIMITACIONES:
        doc.add_paragraph(texto, style="List Bullet")

    salida = Path(args.salida)
    salida.parent.mkdir(parents=True, exist_ok=True)
    doc.save(salida)

    print(f"Escrito: {salida}")
    print("Reparto: " + " · ".join(
        f"{c} {RESPUESTAS[c]}: {reparto.get(c, 0)}" for c in sorted(RESPUESTAS)))
    faltan = len(preguntas) - len(lista)
    if faltan:
        print(f"AVISO: {faltan} preguntas sin entrada en respuestas.json")
    return 0


if __name__ == "__main__":
    sys.exit(main())
